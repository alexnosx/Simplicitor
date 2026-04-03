# simplicitor/app/generators/word_generator.py
# Phase 3: Word document generator
from pathlib import Path

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn


class WordGenerator:
    """Generates .docx files from parsed LLM output (Phase 3)."""

    def generate(self, parsed: dict, output_path: Path) -> Path:
        """Write a .docx file from a parsed Word response dict.

        Args:
            parsed: Validated dict from parse_word_response(). Expected keys:
                    ``title`` (str) and ``sections`` (list of section dicts).
            output_path: Full file path (including filename) to write.

        Returns:
            output_path as a Path.

        Raises:
            OSError: On disk write failure.
        """
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)

        doc = Document()

        # Set default font on the Normal style
        normal_style = doc.styles["Normal"]
        normal_style.font.name = "Calibri"
        normal_style.font.size = Pt(11)

        # Document title as Heading 1
        doc.add_heading(parsed.get("title", ""), level=1)

        for section in parsed.get("sections", []):
            heading = section.get("heading", "")
            if heading:
                doc.add_heading(heading, level=2)

            content = section.get("content", "")
            section_type = section.get("type", "text")

            if section_type == "text":
                self._add_text_section(doc, content)
            elif section_type == "list":
                self._add_list_section(doc, content)
            elif section_type == "table":
                self._add_table_section(doc, content)

        doc.save(str(output_path))
        return output_path

    def _add_text_section(self, doc: Document, content: str) -> None:
        """Add content split on double newlines as Normal paragraphs."""
        chunks = content.split("\n\n")
        for chunk in chunks:
            chunk = chunk.strip()
            if chunk:
                doc.add_paragraph(chunk)

    def _add_list_section(self, doc: Document, content: str) -> None:
        """Add each non-empty line of content as a List Bullet paragraph."""
        # Support both single and double newline separators
        lines = content.replace("\n\n", "\n").split("\n")
        for line in lines:
            line = line.strip()
            if line:
                doc.add_paragraph(line, style="List Bullet")

    def _add_table_section(self, doc: Document, content: str) -> None:
        """Parse tab- or pipe-separated content and add a table.

        The first line of ``content`` is treated as headers; subsequent lines
        are data rows.  Both ``\\t`` and ``|`` are accepted as column separators.
        Leading/trailing ``|`` characters are stripped from pipe-separated lines.
        """
        lines = [ln for ln in content.split("\n") if ln.strip()]
        if not lines:
            return

        def split_line(line: str) -> list[str]:
            """Split a single line on tab or pipe and strip each cell."""
            if "|" in line:
                # Strip leading/trailing pipes, then split
                line = line.strip().strip("|")
                return [cell.strip() for cell in line.split("|")]
            return [cell.strip() for cell in line.split("\t")]

        headers = split_line(lines[0])
        data_rows = [split_line(ln) for ln in lines[1:]]

        num_cols = max(len(headers), max((len(r) for r in data_rows), default=0))
        num_rows = 1 + len(data_rows)

        table = doc.add_table(rows=num_rows, cols=num_cols)
        table.style = "Table Grid"

        # Write header row (bold)
        header_row = table.rows[0]
        for col_idx, header_text in enumerate(headers):
            cell = header_row.cells[col_idx]
            paragraph = cell.paragraphs[0]
            run = paragraph.add_run(header_text)
            run.bold = True

        # Write data rows
        for row_idx, row_data in enumerate(data_rows):
            table_row = table.rows[row_idx + 1]
            for col_idx, cell_text in enumerate(row_data):
                if col_idx < num_cols:
                    table_row.cells[col_idx].text = cell_text
